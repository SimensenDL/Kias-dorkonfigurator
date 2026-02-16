"""
Eksporterer ordretekst for alle dører i dørlisten til Word-dokument (.docx).
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..models.production_list import ProductionList
from ..utils.ordretekst import generer_ordretekst


def export_ordretekst_docx(prod_list: ProductionList, filepath: str) -> None:
    """Genererer Word-dokument med ordretekst for alle dører.

    Args:
        prod_list: Produksjonslisten med dører
        filepath: Filsti for .docx-filen
    """
    doc = Document()

    # Overskrift
    heading = doc.add_heading("TEKSTER FOR GENERELLE KUNDER", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, door in enumerate(prod_list.doors):
        linjer = generer_ordretekst(door.params)
        if not linjer:
            continue

        # Tittel (indeks 0) — fet
        tittel_para = doc.add_paragraph()
        tittel_run = tittel_para.add_run(linjer[0])
        tittel_run.bold = True
        tittel_run.font.size = Pt(12)

        # "Produktdetaljer:" (indeks 1) — fet, understreket
        if len(linjer) > 1:
            detaljer_para = doc.add_paragraph()
            detaljer_run = detaljer_para.add_run(linjer[1])
            detaljer_run.bold = True
            detaljer_run.underline = True

        # Kulepunkt-linjer (indeks 2+)
        for linje in linjer[2:]:
            doc.add_paragraph(linje, style='List Bullet')

        # Mellomrom mellom dører (unntatt siste)
        if i < len(prod_list.doors) - 1:
            doc.add_paragraph()

    doc.save(filepath)
